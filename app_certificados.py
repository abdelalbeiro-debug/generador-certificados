# =================================================================
# SOFTWARE: Generador de Certificados Auditoría PRO
# DESARROLLADOR: Abdel Areiza
# VERSIÓN: 1.0
# DESCRIPCIÓN: Automatización de certificados contables en Word
# =================================================================

import pandas as pd
import streamlit as st
from docx import Document
from io import BytesIO
import re
import zipfile
from docx.shared import Pt

def limpiar_monto(valor):
    """Convierte textos contables con puntos y comas a números reales."""
    if pd.isna(valor) or valor == "": return 0.0
    s = str(valor).strip()
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    s = re.sub(r'[^\d\.-]', '', s)
    try:
        return float(s)
    except:
        return 0.0

def crear_word(fila_datos, nombre_t, nit_t, ano_t, rubros_disponibles, mapa_tipos):
    doc = Document()
    doc.core_properties.author = "Abdel Areiza"
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11) # Bajamos un punto para que quepa todo el formulario

    # 1. Encabezado
    p1 = doc.add_paragraph()
    p1.alignment = 1
    p1.add_run(f"{nombre_t}\n").bold = True
    p1.add_run(f"NIT {nit_t}").bold = True
    
    doc.add_paragraph("\nCERTIFICA QUE\n").alignment = 1
    
    # 2. Cuerpo
    p3 = doc.add_paragraph()
    p3.add_run("La sociedad ").bold = False
    p3.add_run("AGROPECUARIA SANTAMARIA S.A.").bold = True
    p3.add_run(", identificada según ").bold = False
    p3.add_run("NIT.830.075.074-8").bold = True
    p3.add_run(", teniendo en cuenta los saldos registrados en los libros de contabilidad con esta entidad; al corte 31 de diciembre de ")
    p3.add_run(f"{ano_t}").bold = True
    p3.add_run(", presenta los siguientes saldos contables:").bold = False
    
    doc.add_paragraph()

    # --- LISTA MAESTRA DE AUDITORÍA ---
    maestro_activos = [
        "Clientes Nacionales", "Clientes Nacionales Provisión", 
        "Anticipos y avances entregados", "Aportes para futura suscripción Acciones",
        "Obligaciones con particulares por cobrar", "Dividendos y/o participaciones",
        "Cuentas por cobrar a socios y accionistas", "Otras cuentas por cobrar",
        "Depósitos por cobrar"
    ]
    
    maestro_pasivos = [
        "Proveedores", "Depósitos recibidos por pagar",
        "Préstamos a particulares por pagar", "Dineros recibidos por anticipado",
        "Depósitos para acciones", "Cuentas por pagar a terceros",
        "Otras cuentas por pagar", "Cuentas en participación por pagar"
    ]

    # 3. Generar Secciones
    for titulo, lista_maestra in [("De Activos", maestro_activos), ("De Pasivos", maestro_pasivos)]:
        p_tit = doc.add_paragraph()
        p_tit.add_run(f"{titulo}:").bold = True
        p_tit.paragraph_format.space_after = Pt(0)
        
        for rubro in lista_maestra:
            # Buscamos si el rubro existe en los datos del Excel
            valor = fila_datos.get(rubro, 0)
            
            p_item = doc.add_paragraph()
            p_item.paragraph_format.space_after = Pt(0)
            p_item.paragraph_format.line_spacing = 1.0
            
            # Formateo: Nombre del rubro + Espacios + Valor o Línea
            if abs(valor) > 0.01:
                texto_valor = f"$ {abs(valor):,.2f}"
            else:
                texto_valor = "$ ________________"
            
            # Usamos tabulaciones o espacios para alinear (Ajuste visual)
            # Nota: En Word profesional es mejor usar tabuladores, aquí simulamos alineación
            p_item.add_run(f"{rubro}:").bold = False
            # Añadimos un espacio flexible (esto se puede mejorar con tabs en Word)
            p_item.add_run(f" {texto_valor}").bold = False

        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # 4. Firma
    doc.add_paragraph(f"\nEsta certificación se expide el día 10 de febrero de 2026.")
    p_firma = doc.add_paragraph("\nAtentamente\n\nNombre: ______________________\n")
    p_firma.add_run("Firma: _______________\n")
    p_firma.add_run("Contador Público\n")
    p_firma.add_run("CC: _________________\n")
    p_firma.add_run("TP: __________")
    
    target = BytesIO()
    doc.save(target)
    return target

# --- INTERFAZ ---
# 1. Configuración obligatoria (SOLO UNA VEZ)
st.set_page_config(page_title="Generador Auditoría", layout="wide")

with st.sidebar:
    st.header("📌 Centro de Ayuda")
    try:
        with open("Manual_Usuario_Certificados_PRO.pdf", "rb") as f:
            st.download_button(
                label="📘 Descargar Manual de Usuario",
                data=f,
                file_name="Manual_Usuario_Certificados_PRO.pdf",
                mime="application/pdf"
            )
    except FileNotFoundError:
        st.warning("⚠️ Manual no disponible en el servidor.")
        
# 2. Diseño Personalizado (Fondo Slate y Header Azul)
st.markdown("""
    <style>
    /* Fondo Slate-100 */
    .stApp { background-color: #f1f5f9; }
    
    /* Header Azul-900 */
    header[data-testid="stHeader"] { background-color: #0f172a; }
    
    /* Título en el Header */
    .main-title {
        background-color: #0f172a; color: white; padding: 1.5rem;
        border-radius: 0 0 1rem 1rem; margin-top: -3.5rem;
        margin-bottom: 2rem; text-align: center; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1);
    }

    /* Tarjetas Blancas Redondeadas 2xl */
    div[data-testid="stVerticalBlock"] > div:has(div.stMarkdown) {
        background-color: white; padding: 2rem;
        border-radius: 1.5rem; border: 1px solid #e2e8f0;
        box-shadow: 0 10px 15px -3px rgba(0,0,0,0.1);
        margin-bottom: 1.5rem;
    }
    </style>
    
    <div class="main-title">
        <h1 style="margin:0; font-size: 1.8rem; font-family: Arial;">📑 Generador de Certificados Auditoría PRO</h1>
    </div>
""", unsafe_allow_html=True)

# 3. El cargador de archivos (YA NO NECESITAS st.title AQUÍ)
archivo = st.file_uploader("Sube tu Excel", type=["xlsx"])
# 2. Si NO han subido nada, mostramos el mensaje azul de bienvenida
if not archivo:
    st.info("👋 ¡Hola Colega! Por favor, sube tu archivo de Excel arriba para empezar a generar los certificados.")

# 3. Si YA subieron el archivo, procesamos todo una sola vez
else:
    df = pd.read_excel(archivo)
    df.columns = df.columns.str.upper().str.strip()
    
    # Lista de columnas obligatorias
    req = ['AÑO', 'NOMBRE CUENTA', 'TERCERO', 'NIT', 'VALOR EN LIBROS', 'TIPO']
    
    # Verificamos si el Excel es el correcto
    if all(c in df.columns for c in req):
        st.success("✅ ¡Base de datos cargada con éxito!")
        
        # Limpieza de datos (Asegúrate de que estas líneas tengan 8 espacios de sangría)
        df['NIT'] = df['NIT'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
        df['VALOR EN LIBROS'] = df['VALOR EN LIBROS'].apply(limpiar_monto)
        df['NOMBRE CUENTA'] = df['NOMBRE CUENTA'].astype(str).str.strip()
        df['TIPO'] = df['TIPO'].astype(str).str.strip()
        
        # Lógica del certificado
        mapa_tipos = pd.Series(df.TIPO.values, index=df['NOMBRE CUENTA']).to_dict()
        resumen = df.groupby(['AÑO', 'NIT', 'TERCERO', 'NOMBRE CUENTA'])['VALOR EN LIBROS'].sum().unstack(fill_value=0).reset_index()
        rubros = [c for c in resumen.columns if c not in ['AÑO', 'NIT', 'TERCERO']]
        
        resumen['ETIQUETA'] = resumen['TERCERO'] + " (" + resumen['NIT'] + ") - Año: " + resumen['AÑO'].astype(str)
        
        # Buscador de personas
# Buscador de personas
        seleccionados = st.multiselect("Selecciona los certificados:", resumen['ETIQUETA'].unique())
        
        if seleccionados:
            if st.button("🚀 Preparar Descargas"):
                # --- NUEVA LÓGICA PARA EL ARCHIVO ZIP ---
                zip_buffer = BytesIO()
                
                with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zip_file:
                    for i, sel in enumerate(seleccionados):
                        fila = resumen[resumen['ETIQUETA'] == sel].iloc[0]
                        doc_final = crear_word(fila, fila['TERCERO'], fila['NIT'], fila['AÑO'], rubros, mapa_tipos)
                        
                        nombre_archivo = f"Certificado_{fila['NIT']}_{fila['AÑO']}.docx"
                        # Guardamos cada certificado dentro del paquete ZIP
                        zip_file.writestr(nombre_archivo, doc_final.getvalue())
                        
                        # Mantenemos los botones individuales abajo por si acaso
                        st.download_button(
                            label=f"📥 Descargar: {sel}",
                            data=doc_final.getvalue(),
                            file_name=nombre_archivo,
                            key=f"dl_{i}_{fila['NIT']}"
                        )
                
                # --- BOTÓN DE DESCARGA TODO EN UNO ---
                st.markdown("---")
                st.success("🎉 ¡Todos los certificados están listos!")
                st.download_button(
                    label="🎁 DESCARGAR TODO EL LOTE (.ZIP)",
                    data=zip_buffer.getvalue(),
                    file_name="Lote_Certificados_Auditoria.zip",
                    mime="application/zip",
                    use_container_width=True
                )
    else:
        # Este error solo sale si el archivo está mal
        st.error(f"⚠️ El Excel no tiene las columnas correctas. Necesito: {req}")

# 4. Los créditos (Al puro final, sin espacios a la izquierda)
st.markdown("---")
st.caption("🛠️ Desarrollado por **Especializacion en contabilidad y audoria en entornos digitales - 2026**")
